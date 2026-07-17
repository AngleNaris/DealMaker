"""
合同生成核心逻辑：金额大写、模板处理、联系人、设置。
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def project_root() -> str:
    """
    资源根目录（officecli、模板所在）：
    - 环境变量 DEALMAKER_ROOT（Tauri 启动后端时注入）
    - 打包后：后端 exe 同级目录
    - 开发：backend/ 上一级
    """
    env = os.environ.get("DEALMAKER_ROOT")
    if env and os.path.isdir(env):
        return env
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def get_app_dir() -> str:
    """可写配置目录所在根：与 project_root 一致（exe 旁 .contract_tool）。"""
    return project_root()


def get_config_dir() -> str:
    return os.path.join(get_app_dir(), ".contract_tool")


def get_officecli_path() -> str:
    """获取 officecli 可执行文件路径。

    优先：DEALMAKER_RUNTIME（主程序解压到 LocalAppData 的依赖目录）
    其次：DEALMAKER_ROOT / 开发仓库 / 旧版同目录布局
    """
    candidates = []
    runtime = os.environ.get("DEALMAKER_RUNTIME")
    if runtime:
        candidates.append(os.path.join(runtime, "officecli.exe"))
    root = project_root()
    candidates.extend(
        [
            os.path.join(root, "officecli.exe"),
            os.path.join(root, "resources", "officecli.exe"),
            os.path.join(root, "_internal", "officecli.exe"),
        ]
    )
    if getattr(sys, "frozen", False):
        meipass = getattr(sys, "_MEIPASS", None)
        if meipass:
            candidates.insert(0, os.path.join(meipass, "officecli.exe"))
        candidates.append(os.path.join(os.path.dirname(sys.executable), "officecli.exe"))
    for p in candidates:
        if os.path.isfile(p):
            return p
    return "officecli"


def amount_to_chinese(amount: float) -> str:
    """将数字金额转换为中文大写金额。"""
    digits = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
    units = ["", "拾", "佰", "仟"]
    big_units = ["", "万", "亿", "兆"]

    if amount == 0:
        return "零元整"
    if amount < 0:
        return "负" + amount_to_chinese(-amount)

    amount = round(amount, 2)
    int_part = int(amount)
    dec_part = round((amount - int_part) * 100)
    result = ""

    if int_part > 0:
        int_str = str(int_part)
        n = len(int_str)
        zero_flag = False
        for i, ch in enumerate(int_str):
            d = int(ch)
            pos = n - i - 1
            unit_idx = pos % 4
            big_idx = pos // 4
            if d == 0:
                zero_flag = True
                if unit_idx == 0 and big_idx > 0:
                    result += big_units[big_idx]
            else:
                if zero_flag:
                    result += "零"
                    zero_flag = False
                result += digits[d] + units[unit_idx]
                if unit_idx == 0 and big_idx > 0:
                    result += big_units[big_idx]
        result += "元"

    if dec_part == 0:
        result += "整"
    else:
        jiao = dec_part // 10
        fen = dec_part % 10
        if jiao > 0:
            result += digits[jiao] + "角"
        elif fen > 0:
            result += "零"
        if fen > 0:
            result += digits[fen] + "分"
    return result


def split_by_ratio(total: float, ratio_percent: float) -> tuple[float, float]:
    """按比例拆分预付款/尾款，保证两者之和等于 total。"""
    prepaid = round(total * ratio_percent / 100.0, 2)
    final = round(total - prepaid, 2)
    return prepaid, final


def auto_fix_final(total: float, prepaid: float) -> tuple[float, float]:
    """策略 A：以预付款为准修正尾款。"""
    prepaid = round(prepaid, 2)
    if prepaid > total:
        prepaid = round(total, 2)
        return prepaid, 0.0
    if prepaid < 0:
        prepaid = 0.0
    return prepaid, round(total - prepaid, 2)


class TemplateProcessor:
    """模板处理器：加载模板、识别占位符、替换内容。"""

    PLACEHOLDER_PATTERN = re.compile(r"%([^%]+)%")

    def __init__(self, template_path: str):
        self.template_path = template_path
        self.doc = None
        self.placeholders: List[str] = []
        self.load_template()

    def load_template(self):
        self.doc = Document(self.template_path)
        self.placeholders = self._find_placeholders()

    def _find_placeholders(self) -> List[str]:
        placeholders = set()

        def extract(text: str):
            text = text.replace("%%", "%\x00")
            for m in self.PLACEHOLDER_PATTERN.findall(text):
                if "\x00" not in m:
                    placeholders.add(m)

        for para in self.doc.paragraphs:
            extract(para.text)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        extract(para.text)
        return sorted(placeholders)

    def _resolve_placeholder(self, key: str) -> str:
        if key in self.placeholders:
            return key
        for ph in self.placeholders:
            if ph.startswith(key):
                return ph
        return key

    def generate(self, replacements: Dict[str, str], output_path: str) -> str:
        shutil.copy2(self.template_path, output_path)
        processed = dict(replacements)

        address = processed.pop("乙方地址", "")
        if address:
            lines: List[str] = []
            remaining = address
            for max_len in (17, 20, 20):
                if not remaining:
                    break
                if len(remaining) <= max_len:
                    lines.append(remaining)
                    remaining = ""
                else:
                    break_point = max_len
                    for i in range(max_len - 1, max_len // 2, -1):
                        if remaining[i] in "，。、；：！？,.;:!? ":
                            break_point = i + 1
                            break
                    lines.append(remaining[:break_point])
                    remaining = remaining[break_point:]
            if remaining:
                if lines:
                    lines[-1] += remaining
                else:
                    lines.append(remaining)
            while len(lines) < 3:
                lines.append("")
            processed["替换的乙方地址第一行最大字数"] = lines[0]
            processed["替换的乙方地址第二行最大字数最大字"] = lines[1]
            processed["替换的乙方地址第三行最大字数最大字"] = lines[2]

        img_path = processed.pop("替换的费用表格图片", "")

        for key in list(processed.keys()):
            if "大写" in key:
                actual_key = self._resolve_placeholder(key)
                v = str(processed[key])
                if not actual_key.endswith("整") and not actual_key.endswith("元整"):
                    if v.endswith("元整"):
                        processed[key] = v[:-2]
                    elif v.endswith("整"):
                        processed[key] = v[:-1]

        commands = []
        for key, value in processed.items():
            actual_key = self._resolve_placeholder(key)
            commands.append(
                {
                    "command": "set",
                    "path": "/",
                    "props": {"find": f"%{actual_key}%", "replace": str(value)},
                }
            )

        if commands:
            batch_file = output_path + ".batch.json"
            with open(batch_file, "w", encoding="utf-8") as f:
                json.dump(commands, f, ensure_ascii=False)
            result = subprocess.run(
                [get_officecli_path(), "batch", output_path, "--input", batch_file],
                capture_output=True,
                text=True,
            )
            try:
                os.remove(batch_file)
            except OSError:
                pass
            if result.returncode != 0:
                err = (result.stderr or result.stdout or "officecli batch 失败").strip()
                raise RuntimeError(err)

        if img_path and os.path.isfile(img_path):
            doc = Document(output_path)
            self._insert_image(doc, img_path)
            doc.save(output_path)

        doc = Document(output_path)
        self._normalize_address_font(doc)
        doc.save(output_path)
        return output_path

    def _content_width_emu(self, doc) -> int:
        """当前节正文可用宽度（页宽 - 左右边距），EMU。"""
        section = doc.sections[0]
        # 若图片所在段落属于其它节，尽量用对应节；默认第一节
        usable = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
        # 略留一点余量，避免贴边溢出
        usable = max(usable - int(Inches(0.05)), int(Inches(1)))
        return usable

    def _reset_paragraph_for_image(self, para):
        """清空占位段并重建为纯居中段落（清除 WPS 段落布局/样式/缩进）。

        模板占位段可能带 Footer 样式、左缩进、framePr（段落布局）或权限标记，
        仅改 alignment 在 WPS 下仍可能偏左。此处等价于：清除段落布局 →
        去掉特殊样式 → 居中 → 再插入图片。
        """
        p = para._p
        # 去掉 pPr 以外的全部子节点（runs / permStart / permEnd 等）
        for child in list(p):
            if child.tag != qn("w:pPr"):
                p.remove(child)
        # 重建干净的 pPr：仅保留居中
        old_pPr = p.find(qn("w:pPr"))
        if old_pPr is not None:
            p.remove(old_pPr)
        pPr = p.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _insert_image(self, doc, img_path: str):
        """插入费用表格图片，宽度自适应页面正文宽度，高度按比例缩放，段落居中。"""
        width = self._content_width_emu(doc)

        def place_in_paragraph(para) -> bool:
            if "替换的费用表格图片" not in para.text:
                return False
            self._reset_paragraph_for_image(para)
            run = para.add_run()
            # 只设 width 时 python-docx 保持原图宽高比（inline，非浮动）
            run.add_picture(img_path, width=width)
            return True

        for para in doc.paragraphs:
            if place_in_paragraph(para):
                return

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if place_in_paragraph(para):
                            return

    def _normalize_address_font(self, doc):
        addr_start = -1
        for i, para in enumerate(doc.paragraphs):
            if "乙方:" in para.text or "乙方：" in para.text:
                addr_start = i + 1
                break
        if addr_start < 0:
            return
        for offset in range(3):
            idx = addr_start + offset
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                for run in para.runs:
                    if run.text.strip():
                        run.font.name = "宋体"
                        run.font.size = Pt(12)


# 联系人仅保存乙方信息（不含合同编号/项目/服务/费用等项目字段）
CONTACT_FIELD_KEYS = (
    "替换的乙方名称",
    "乙方名称",  # 兼容旧键
    "乙方银行账号",
    "乙方银行开户行",
    "替换的乙方代表名称",
    "替换的乙方代表电话",
    "替换的乙方代表邮箱",
    "乙方地址",
)


def pick_contact_fields(data: Dict) -> Dict:
    """从完整表单中提取联系人字段。"""
    out: Dict = {}
    for k in CONTACT_FIELD_KEYS:
        if k in data and data[k] is not None:
            out[k] = data[k]
    name = (out.get("替换的乙方名称") or out.get("乙方名称") or "").strip()
    if name:
        out["替换的乙方名称"] = name
        out.pop("乙方名称", None)
    return out


class ContactManager:
    def __init__(self, config_dir: Optional[str] = None):
        self.config_dir = config_dir or get_config_dir()
        self.contacts_file = os.path.join(self.config_dir, "contacts.json")
        self.contacts: List[Dict] = []
        self.load_contacts()

    def load_contacts(self):
        if os.path.exists(self.contacts_file):
            try:
                with open(self.contacts_file, "r", encoding="utf-8") as f:
                    raw = json.load(f)
                # 读入时规范化为仅乙方字段（旧数据里可能混有项目信息）
                self.contacts = [pick_contact_fields(c) for c in (raw or []) if isinstance(c, dict)]
            except Exception:
                self.contacts = []
        else:
            self.contacts = []

    def save_contacts(self):
        os.makedirs(self.config_dir, exist_ok=True)
        with open(self.contacts_file, "w", encoding="utf-8") as f:
            json.dump(self.contacts, f, ensure_ascii=False, indent=2)

    def _get_company_name(self, contact: Dict) -> str:
        return contact.get("替换的乙方名称", "") or contact.get("乙方名称", "")

    def add_contact(self, contact: Dict):
        contact = pick_contact_fields(contact or {})
        name = self._get_company_name(contact)
        if not name:
            return
        for i, c in enumerate(self.contacts):
            if self._get_company_name(c) == name:
                self.contacts[i] = contact
                self.save_contacts()
                return
        self.contacts.append(contact)
        self.save_contacts()

    def delete_contact(self, name: str):
        self.contacts = [c for c in self.contacts if self._get_company_name(c) != name]
        self.save_contacts()

    def get_contact(self, name: str) -> Optional[Dict]:
        for c in self.contacts:
            if self._get_company_name(c) == name:
                return c
        return None

    def get_names(self) -> List[str]:
        return [self._get_company_name(c) for c in self.contacts if self._get_company_name(c)]


def load_settings() -> Dict:
    path = os.path.join(get_config_dir(), "settings.json")
    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_settings(settings: Dict) -> None:
    os.makedirs(get_config_dir(), exist_ok=True)
    path = os.path.join(get_config_dir(), "settings.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(settings, f, ensure_ascii=False, indent=2)


# ---------- 历史项目（应用内数据，与导出文件无关）----------

def make_project_key(contract_no: str, project_name: str) -> str:
    """合同编号 + 项目名称 组成唯一键。"""
    return f"{(contract_no or '').strip()}||{(project_name or '').strip()}"


class ProjectStore:
    """合同/报价填写快照存储。"""

    def __init__(self, config_dir: Optional[str] = None):
        self.config_dir = config_dir or get_config_dir()
        self.path = os.path.join(self.config_dir, "projects.json")
        self.projects: List[Dict] = []
        self.load()

    def load(self) -> None:
        if os.path.exists(self.path):
            try:
                with open(self.path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                self.projects = data if isinstance(data, list) else data.get("projects", [])
            except Exception:
                self.projects = []
        else:
            self.projects = []

    def save(self) -> None:
        os.makedirs(self.config_dir, exist_ok=True)
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(self.projects, f, ensure_ascii=False, indent=2)

    def list_summaries(self) -> List[Dict]:
        """列表用摘要，按更新时间倒序。"""
        items = sorted(
            self.projects,
            key=lambda p: p.get("updated_at") or "",
            reverse=True,
        )
        out = []
        for p in items:
            out.append(
                {
                    "id": p.get("id"),
                    "key": p.get("key"),
                    "contract_no": p.get("contract_no", ""),
                    "project_name": p.get("project_name", ""),
                    "party_b": p.get("party_b", ""),
                    "updated_at": p.get("updated_at", ""),
                    "label": self._label(p),
                }
            )
        return out

    def _label(self, p: Dict) -> str:
        no = (p.get("contract_no") or "").strip() or "无编号"
        name = (p.get("project_name") or "").strip() or "未命名项目"
        return f"{no} · {name}"

    def get(self, project_id: str) -> Optional[Dict]:
        for p in self.projects:
            if p.get("id") == project_id:
                return p
        return None

    def get_by_key(self, key: str) -> Optional[Dict]:
        for p in self.projects:
            if p.get("key") == key:
                return p
        return None

    def upsert(self, snapshot: Dict) -> Dict:
        """
        按 合同编号+项目名称 键保存。
        键已存在 → 更新原项目；否则新建。
        snapshot 字段见 CLI。
        """
        import uuid
        from datetime import datetime, timezone

        contract_no = (snapshot.get("contract_no") or "").strip()
        project_name = (snapshot.get("project_name") or "").strip()
        if not contract_no and not project_name:
            raise ValueError("请至少填写合同编号或项目名称后再保存项目")

        key = make_project_key(contract_no, project_name)
        now = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M:%S")
        form = snapshot.get("form") or {}
        party_b = (
            form.get("替换的乙方名称")
            or form.get("乙方名称")
            or snapshot.get("party_b")
            or ""
        )

        existing = self.get_by_key(key)
        if existing:
            existing.update(
                {
                    "contract_no": contract_no,
                    "project_name": project_name,
                    "party_b": party_b,
                    "form": form,
                    "ratio": snapshot.get("ratio", 50),
                    "quote": snapshot.get("quote"),
                    "template_path": snapshot.get("template_path") or "",
                    "output_dir": snapshot.get("output_dir") or "",
                    "output_name": snapshot.get("output_name") or "",
                    "updated_at": now,
                }
            )
            self.save()
            return {"action": "updated", "project": existing}

        proj = {
            "id": str(uuid.uuid4()),
            "key": key,
            "contract_no": contract_no,
            "project_name": project_name,
            "party_b": party_b,
            "form": form,
            "ratio": snapshot.get("ratio", 50),
            "quote": snapshot.get("quote"),
            "template_path": snapshot.get("template_path") or "",
            "output_dir": snapshot.get("output_dir") or "",
            "output_name": snapshot.get("output_name") or "",
            "created_at": now,
            "updated_at": now,
        }
        self.projects.append(proj)
        self.save()
        return {"action": "created", "project": proj}

    def delete(self, project_id: str) -> bool:
        before = len(self.projects)
        self.projects = [p for p in self.projects if p.get("id") != project_id]
        if len(self.projects) < before:
            self.save()
            return True
        return False



def default_template_path() -> Optional[str]:
    root = get_app_dir()
    candidates = [
        os.path.join(root, "_合同模板.docx"),
        os.path.join(root, "resources", "_合同模板.docx"),
        os.path.join(root, "contract_template.docx"),
    ]
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        candidates.insert(0, os.path.join(meipass, "_合同模板.docx"))
    for p in candidates:
        if os.path.isfile(p):
            return p
    return None


def build_output_path(
    template_path: str,
    data: Dict[str, str],
    output_dir: str = "",
    output_name: str = "",
    ext: str = "docx",
) -> str:
    out_dir = output_dir or os.path.dirname(template_path) or get_app_dir()
    name = output_name
    if not name:
        contract_no = data.get("替换的合同编号", "")
        project_name = data.get("替换的项目名称", "")
        company = data.get("替换的乙方名称", "")
        parts = []
        if contract_no:
            parts.append(contract_no)
        if project_name and company:
            parts.append(f"{project_name}（{company}）")
        elif project_name:
            parts.append(project_name)
        elif company:
            parts.append(company)
        name = "_".join(parts) if parts else "合同"
    return os.path.join(out_dir, f"{name}.{ext}")


def _com_progid_available(progid: str) -> bool:
    if sys.platform != "win32":
        return False
    try:
        import winreg

        winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, progid).Close()
        return True
    except OSError:
        return False


def _wps_com_available() -> bool:
    """WPS 文字 COM（Kwps.Application）。"""
    return _com_progid_available("Kwps.Application") or _com_progid_available(
        "KWPS.Application"
    )


def _word_com_available() -> bool:
    return _com_progid_available("Word.Application")


def pdf_engines_status() -> Dict[str, bool]:
    return {
        "wps": _wps_com_available(),
        "word": _word_com_available(),
    }


def docx_to_pdf(docx_path: str, pdf_path: Optional[str] = None) -> Dict[str, str]:
    """
    将 DOCX 转为 PDF（高保真）。

    优先级：WPS → Microsoft Word。
    均不可用时提示安装，不使用 LibreOffice。
    返回 {path, engine}，engine 为 wps | word。
    """
    docx_path = os.path.abspath(docx_path)
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"找不到文档: {docx_path}")

    if not pdf_path:
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    else:
        pdf_path = os.path.abspath(pdf_path)

    os.makedirs(os.path.dirname(pdf_path) or ".", exist_ok=True)
    if os.path.exists(pdf_path):
        try:
            os.remove(pdf_path)
        except OSError:
            pass

    errors: List[str] = []
    has_wps = _wps_com_available()
    has_word = _word_com_available()

    if has_wps:
        try:
            return _docx_to_pdf_com(docx_path, pdf_path, engine="wps")
        except Exception as e:
            errors.append(f"WPS: {e}")

    if has_word:
        try:
            return _docx_to_pdf_com(docx_path, pdf_path, engine="word")
        except Exception as e:
            errors.append(f"Word: {e}")

    if not has_wps and not has_word:
        raise RuntimeError(
            "无法导出 PDF：未检测到 WPS 文字 或 Microsoft Word。\n"
            "请安装其一后重试：\n"
            "  · WPS Office（推荐，本机 COM：Kwps.Application）\n"
            "  · Microsoft Word 桌面版\n"
            "安装完成后重启 DealMaker。"
        )

    detail = "\n".join(errors) if errors else "未知错误"
    raise RuntimeError(
        "PDF 导出失败（已尝试本机可用的办公软件）：\n"
        f"{detail}\n"
        "请确认 WPS/Word 可正常打开文档，关闭其弹窗后重试。"
    )


def _docx_to_pdf_com(docx_path: str, pdf_path: str, engine: str) -> Dict[str, str]:
    """
    通过 COM ExportAsFixedFormat 导出 PDF。
    engine: wps | word
    """
    if engine == "wps":
        progid = "Kwps.Application"
        label = "WPS"
    elif engine == "word":
        progid = "Word.Application"
        label = "Word"
    else:
        raise ValueError(f"未知引擎: {engine}")

    def esc(p: str) -> str:
        return p.replace("'", "''")

    # PowerShell 调用 COM，避免依赖 pywin32
    ps = f"""
$ErrorActionPreference = 'Stop'
$docx = '{esc(docx_path)}'
$pdf = '{esc(pdf_path)}'
$app = $null
$doc = $null
try {{
  $app = New-Object -ComObject {progid}
  try {{ $app.Visible = $false }} catch {{}}
  try {{ $app.DisplayAlerts = 0 }} catch {{}}
  $doc = $app.Documents.Open($docx)
  $doc.ExportAsFixedFormat($pdf, 17)
  if (-not (Test-Path -LiteralPath $pdf)) {{ throw '{label} 未生成 PDF 文件' }}
}} finally {{
  if ($null -ne $doc) {{
    try {{ $doc.Close($false) }} catch {{ try {{ $doc.Close() }} catch {{}} }}
  }}
  if ($null -ne $app) {{
    try {{ $app.Quit() }} catch {{}}
    try {{ [void][System.Runtime.InteropServices.Marshal]::FinalReleaseComObject($app) }} catch {{}}
  }}
  [GC]::Collect(); [GC]::WaitForPendingFinalizers()
}}
"""
    result = subprocess.run(
        [
            "powershell",
            "-NoProfile",
            "-NonInteractive",
            "-WindowStyle",
            "Hidden",
            "-Command",
            ps,
        ],
        capture_output=True,
        text=True,
        timeout=180,
    )
    if result.returncode != 0 or not os.path.isfile(pdf_path):
        err = (result.stderr or result.stdout or f"{label} 导出失败").strip()
        raise RuntimeError(err)
    return {"path": pdf_path, "engine": engine}
