"""
合同模板生成工具
功能：读取Word模板，填入数据，输出DOCX和PDF
作者：哈雷酱 (傲娇大小姐工程师)
"""

import sys
import os
import json
import re
import subprocess
from pathlib import Path
from typing import Dict, List, Optional

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QFileDialog, QScrollArea,
    QGroupBox, QComboBox, QMessageBox, QTabWidget, QFormLayout,
    QSplitter, QFrame, QSizePolicy, QListWidget, QAbstractItemView
)
from PyQt6.QtCore import Qt, QSize
from PyQt6.QtGui import QFont, QIcon

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy


# ============================================================
# officecli 路径（兼容打包模式）
# ============================================================

def get_officecli_path():
    """获取 officecli 可执行文件路径，兼容 PyInstaller 打包和开发模式"""
    if getattr(sys, 'frozen', False):
        # onefile 模式：从临时解压目录查找
        meipass = getattr(sys, '_MEIPASS', None)
        if meipass:
            p = os.path.join(meipass, 'officecli.exe')
            if os.path.exists(p):
                return p
        # onefolder 模式
        base = os.path.dirname(sys.executable)
        candidates = [
            os.path.join(base, '_internal', 'officecli.exe'),
            os.path.join(base, 'officecli.exe'),
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
    base = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base, 'officecli.exe')
    if os.path.exists(path):
        return path
    return 'officecli'


# ============================================================
# 可拖拽的输入框
# ============================================================

class DropLineEdit(QLineEdit):
    """支持拖拽文件到输入框"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            self.setText(path)


# ============================================================
# 金额大写转换
# ============================================================

def amount_to_chinese(amount: float) -> str:
    """将数字金额转换为中文大写金额"""
    digits = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    units = ['', '拾', '佰', '仟']
    big_units = ['', '万', '亿', '兆']

    if amount == 0:
        return '零元整'

    # 处理负数
    if amount < 0:
        return '负' + amount_to_chinese(-amount)

    # 分离整数和小数部分
    amount = round(amount, 2)
    int_part = int(amount)
    dec_part = round((amount - int_part) * 100)

    result = ''

    # 处理整数部分
    if int_part > 0:
        int_str = str(int_part)
        n = len(int_str)
        zero_flag = False

        for i, ch in enumerate(int_str):
            d = int(ch)
            pos = n - i - 1  # 从右往左的位置
            unit_idx = pos % 4
            big_idx = pos // 4

            if d == 0:
                zero_flag = True
                # 如果是万/亿位，需要加单位
                if unit_idx == 0 and big_idx > 0:
                    result += big_units[big_idx]
            else:
                if zero_flag:
                    result += '零'
                    zero_flag = False
                result += digits[d] + units[unit_idx]
                # 如果是万/亿位边界
                if unit_idx == 0 and big_idx > 0:
                    result += big_units[big_idx]

        result += '元'

    # 处理小数部分
    if dec_part == 0:
        result += '整'
    else:
        jiao = dec_part // 10
        fen = dec_part % 10
        if jiao > 0:
            result += digits[jiao] + '角'
        elif fen > 0:
            result += '零'
        if fen > 0:
            result += digits[fen] + '分'

    return result


# ============================================================
# 模板处理
# ============================================================

class TemplateProcessor:
    """模板处理器，负责加载模板、识别占位符、替换内容"""

    PLACEHOLDER_PATTERN = re.compile(r'%([^%]+)%')

    def __init__(self, template_path: str):
        self.template_path = template_path
        self.doc = None
        self.placeholders: List[str] = []
        self.load_template()

    def load_template(self):
        """加载模板并识别占位符"""
        self.doc = Document(self.template_path)
        self.placeholders = self._find_placeholders()

    def _find_placeholders(self) -> List[str]:
        """遍历文档找到所有占位符。处理 %% 边界情况：
        Word run 拼接中 %% 的第二个 % 会被错误匹配为垃圾占位符的起始符，
        同时消耗掉下一个有效占位符的 % 起始符。将 %% 替换为 %\x00 解决。"""
        placeholders = set()

        def extract(text):
            text = text.replace('%%', '%\x00')
            for m in self.PLACEHOLDER_PATTERN.findall(text):
                if '\x00' not in m:
                    placeholders.add(m)

        for para in self.doc.paragraphs:
            extract(para.text)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        extract(para.text)

        return sorted(list(placeholders))

    def _resolve_placeholder(self, key: str) -> str:
        """将表单 key 映射为模板中实际存在的占位符名。
        模板中大写金额占位符可能带后缀（整/元整），如：
        - 表单 key: '替换的总费用大写' → 模板: '替换的总费用大写整'
        - 表单 key: '替换的预付款大写' → 模板: '替换的预付款大写元整'
        """
        if key in self.placeholders:
            return key
        for ph in self.placeholders:
            if ph.startswith(key):
                return ph
        return key

    def generate(self, replacements: Dict[str, str], output_path: str):
        """生成替换后的文档（officecli 负责文本替换，python-docx 负责图片和排版）"""
        # 1. 复制模板到输出路径
        import shutil
        shutil.copy2(self.template_path, output_path)

        # 2. 准备替换数据
        processed = dict(replacements)

        # 处理地址分行
        address = processed.pop('乙方地址', '')
        if address:
            max_len_1, max_len_2, max_len_3 = 17, 20, 20
            lines = []
            remaining = address
            for max_len in [max_len_1, max_len_2, max_len_3]:
                if remaining:
                    if len(remaining) <= max_len:
                        lines.append(remaining)
                        remaining = ''
                    else:
                        break_point = max_len
                        for i in range(max_len - 1, max_len // 2, -1):
                            if remaining[i] in '，。、；：！？,.;:!? ':
                                break_point = i + 1
                                break
                        lines.append(remaining[:break_point])
                        remaining = remaining[break_point:]
            if remaining:
                if lines:
                    lines[-1] += remaining
                else:
                    lines.append(remaining)
            while len(lines) < 3:
                lines.append('')
            processed['替换的乙方地址第一行最大字数'] = lines[0]
            processed['替换的乙方地址第二行最大字数最大字'] = lines[1]
            processed['替换的乙方地址第三行最大字数最大字'] = lines[2]

        # 图片占位符不参与文本替换（后续用 python-docx 插入）
        img_path = processed.pop('替换的费用表格图片', '')

        # 处理大写金额：模板占位符外已自带 整/元整，替换值末尾不再重复
        for key in list(processed.keys()):
            if '大写' in key:
                actual_key = self._resolve_placeholder(key)
                v = str(processed[key])
                if not actual_key.endswith('整') and not actual_key.endswith('元整'):
                    if v.endswith('元整'):
                        processed[key] = v[:-2]
                    elif v.endswith('整'):
                        processed[key] = v[:-1]

        # 3. 构建 officecli batch 命令
        commands = []
        for key, value in processed.items():
            actual_key = self._resolve_placeholder(key)
            commands.append({
                'command': 'set',
                'path': '/',
                'props': {'find': f'%{actual_key}%', 'replace': str(value)}
            })

        if commands:
            batch_file = output_path + '.batch.json'
            with open(batch_file, 'w', encoding='utf-8') as f:
                json.dump(commands, f, ensure_ascii=False)
            subprocess.run(
                [get_officecli_path(), 'batch', output_path, '--input', batch_file],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )
            os.remove(batch_file)

        # 4. 图片插入（python-docx）
        if img_path and os.path.isfile(img_path):
            doc = Document(output_path)
            self._insert_image_before_replace(doc, img_path)
            doc.save(output_path)

        # 5. 地址排版清理（python-docx）
        doc = Document(output_path)
        self._normalize_address_font(doc)
        doc.save(output_path)

        return output_path

    def _insert_image_before_replace(self, doc, img_path: str):
        """在文本替换之前插入图片（查找占位符段落）"""
        for para in doc.paragraphs:
            text = para.text
            if '替换的费用表格图片' in text:
                for run in para.runs:
                    run.text = ''
                run = para.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return

    def _normalize_address_font(self, doc):
        """统一地址行字体为宋体 12pt"""
        addr_start = -1
        for i, para in enumerate(doc.paragraphs):
            if '乙方:' in para.text or '乙方：' in para.text:
                addr_start = i + 1
                break
        if addr_start < 0:
            return
        for offset in range(3):
            idx = addr_start + offset
            if idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                for run in para.runs:
                    if run.text.strip():
                        run.font.name = '宋体'
                        run.font.size = Pt(12)


# ============================================================
# 联系人管理
# ============================================================

class ContactManager:
    """联系人/乙方信息管理器"""

    def __init__(self, config_dir: str):
        self.config_dir = config_dir
        self.contacts_file = os.path.join(config_dir, 'contacts.json')
        self.contacts: List[Dict] = []
        self.load_contacts()

    def load_contacts(self):
        """加载联系人数据"""
        if os.path.exists(self.contacts_file):
            try:
                with open(self.contacts_file, 'r', encoding='utf-8') as f:
                    self.contacts = json.load(f)
            except:
                self.contacts = []
        else:
            self.contacts = []

    def save_contacts(self):
        """保存联系人数据"""
        os.makedirs(self.config_dir, exist_ok=True)
        with open(self.contacts_file, 'w', encoding='utf-8') as f:
            json.dump(self.contacts, f, ensure_ascii=False, indent=2)

    def _get_company_name(self, contact: Dict) -> str:
        """获取联系人的公司名称（兼容两种key格式）"""
        return contact.get('替换的乙方名称', '') or contact.get('乙方名称', '')

    def add_contact(self, contact: Dict):
        """添加或更新联系人"""
        name = self._get_company_name(contact)
        if not name:
            return
        for i, c in enumerate(self.contacts):
            if self._get_company_name(c) == name:
                self.contacts[i] = contact
                self.save_contacts()
                return
        self.contacts.append(contact)
        self.save_contacts()

    def delete_contact(self, name: str):
        """删除联系人"""
        self.contacts = [c for c in self.contacts if self._get_company_name(c) != name]
        self.save_contacts()

    def get_contact(self, name: str) -> Optional[Dict]:
        """获取联系人信息"""
        for c in self.contacts:
            if self._get_company_name(c) == name:
                return c
        return None

    def get_names(self) -> List[str]:
        """获取所有联系人公司名称"""
        return [self._get_company_name(c) for c in self.contacts if self._get_company_name(c)]


# ============================================================
# GUI 主窗口
# ============================================================

class MainWindow(QMainWindow):
    """主窗口"""

    def __init__(self):
        super().__init__()

        # 初始化变量
        self.template_processor: Optional[TemplateProcessor] = None
        # 可写目录：打包后放在 exe 同级，开发模式放在脚本同级
        if getattr(sys, 'frozen', False):
            self._app_dir = os.path.dirname(sys.executable)
        else:
            self._app_dir = os.path.dirname(os.path.abspath(__file__))
        self.config_dir = os.path.join(self._app_dir, '.contract_tool')
        self.contact_manager = ContactManager(self.config_dir)
        self.input_fields: Dict[str, QLineEdit] = {}
        self.settings_file = os.path.join(self.config_dir, 'settings.json')

        # 窗口设置
        self.setWindowTitle("DealMaker")
        self.setMinimumSize(900, 650)
        self.resize(1050, 750)

        # 初始化UI
        self.init_ui()

        # 恢复上次设置
        self._load_settings()

        # 尝试加载默认模板
        if not self.template_path.text():
            # 打包后从解压目录找，开发模式从脚本目录找
            candidates = []
            meipass = getattr(sys, '_MEIPASS', None)
            if meipass:
                candidates.append(os.path.join(meipass, '_合同模板.docx'))
            candidates.append(os.path.join(self._app_dir, '_合同模板.docx'))
            for p in candidates:
                if os.path.exists(p):
                    self.template_path.setText(p)
                    self.load_template()
                    break

    def init_ui(self):
        """初始化界面"""
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)

        # ===== 品牌标题栏 =====
        title_widget = QWidget()
        title_layout = QHBoxLayout(title_widget)
        title_layout.setContentsMargins(0, 0, 0, 4)

        app_title = QLabel("DealMaker")
        app_title.setObjectName("appTitle")
        title_layout.addWidget(app_title)

        title_layout.addStretch()

        author_label = QLabel("@繁星之子卡萨蒂亚")
        author_label.setObjectName("authorLabel")
        title_layout.addWidget(author_label)

        layout.addWidget(title_widget)

        # ===== 模板选择区域 =====
        template_group = QGroupBox("模板文件")
        template_layout = QHBoxLayout(template_group)

        self.template_path = QLineEdit()
        self.template_path.setPlaceholderText("选择合同模板文件...")
        self.template_path.setReadOnly(True)
        template_layout.addWidget(self.template_path)

        btn_select = QPushButton("选择文件")
        btn_select.clicked.connect(self.select_template)
        template_layout.addWidget(btn_select)

        layout.addWidget(template_group)

        # ===== 主体区域（使用分割器，占满可用高度）=====
        splitter = QSplitter(Qt.Orientation.Horizontal)

        # 左侧：联系人管理
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)

        contact_group = QGroupBox("联系人管理")
        contact_layout = QVBoxLayout(contact_group)

        # 联系人选择
        contact_select_layout = QHBoxLayout()
        self.contact_combo = QComboBox()
        self.contact_combo.setPlaceholderText("-- 选择已有联系人 --")
        self.contact_combo.currentTextChanged.connect(self.on_contact_selected)
        contact_select_layout.addWidget(self.contact_combo)

        btn_load_contact = QPushButton("加载")
        btn_load_contact.clicked.connect(self.load_selected_contact)
        contact_select_layout.addWidget(btn_load_contact)

        contact_layout.addLayout(contact_select_layout)

        # 联系人操作按钮
        contact_btn_layout = QHBoxLayout()
        btn_save_contact = QPushButton("保存当前为联系人")
        btn_save_contact.clicked.connect(self.save_current_as_contact)
        contact_btn_layout.addWidget(btn_save_contact)

        btn_new_contact = QPushButton("新建")
        btn_new_contact.clicked.connect(self.clear_form_for_new_contact)
        contact_btn_layout.addWidget(btn_new_contact)

        btn_delete_contact = QPushButton("删除选中")
        btn_delete_contact.clicked.connect(self.delete_selected_contact)
        contact_btn_layout.addWidget(btn_delete_contact)

        contact_layout.addLayout(contact_btn_layout)
        left_layout.addWidget(contact_group)

        # 联系人列表
        self.contact_list = QListWidget()
        self.contact_list.itemClicked.connect(self.on_contact_list_clicked)
        left_layout.addWidget(self.contact_list, 1)

        # 右侧：表单区域（占满可用空间）
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)

        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        scroll_widget = QWidget()
        self.form_layout = QVBoxLayout(scroll_widget)

        self.form_placeholder = QLabel("加载模板后将自动生成表单...")
        self.form_layout.addWidget(self.form_placeholder)
        self.form_layout.addStretch()

        scroll_area.setWidget(scroll_widget)
        right_layout.addWidget(scroll_area)

        # 添加到分割器
        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        splitter.setSizes([220, 680])
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)

        layout.addWidget(splitter, 1)  # stretch factor 1 = 占满所有可用高度

        # ===== 输出设置区域 =====
        output_group = QGroupBox("输出设置")
        output_layout = QHBoxLayout(output_group)

        output_layout.addWidget(QLabel("输出目录:"))
        self.output_dir = QLineEdit()
        self.output_dir.setPlaceholderText("默认与模板同目录")
        output_layout.addWidget(self.output_dir)

        btn_output_dir = QPushButton("选择")
        btn_output_dir.clicked.connect(self.select_output_dir)
        output_layout.addWidget(btn_output_dir)

        output_layout.addWidget(QLabel("文件名:"))
        self.output_name = QLineEdit()
        self.output_name.setPlaceholderText("合同编号_乙方名称")
        self.output_name.setFixedWidth(200)
        output_layout.addWidget(self.output_name)

        layout.addWidget(output_group)

        # ===== 生成按钮区域 =====
        btn_layout = QHBoxLayout()

        btn_generate = QPushButton("生成 DOCX")
        btn_generate.setObjectName("btnGenDocx")
        btn_generate.clicked.connect(self.generate)
        btn_layout.addWidget(btn_generate)

        layout.addLayout(btn_layout)

        # 初始化联系人下拉框
        self.refresh_contact_combo()

    def select_template(self):
        """选择模板文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择合同模板", "",
            "Word文档 (*.docx);;所有文件 (*.*)"
        )
        if file_path:
            self.template_path.setText(file_path)
            self.load_template()
            self._save_settings()

    def load_template(self):
        """加载模板并生成表单"""
        template_path = self.template_path.text()
        if not template_path or not os.path.exists(template_path):
            return

        try:
            self.template_processor = TemplateProcessor(template_path)
            self.generate_form()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"加载模板失败:\n{str(e)}")

    def generate_form(self):
        """根据占位符动态生成表单"""
        # 清空现有表单
        while self.form_layout.count():
            item = self.form_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.input_fields.clear()

        if not self.template_processor:
            return

        # 按合同内容顺序定义分组
        groups = [
            ("合同基本信息", [
                ('合同编号', '请输入合同编号'),
                ('项目名称', '请输入项目名称'),
            ]),
            ("乙方基本信息", [
                ('乙方名称', '请输入乙方公司名称'),
            ]),
            ("服务内容", [
                ('服务内容', '请输入服务内容描述'),
                ('交付格式', '请输入交付物格式'),
                ('交付时间', '请输入交付时间'),
            ]),
            ("费用信息", [
                ('总费用', '输入数字，自动转换大写'),
                ('总费用大写', '自动转换，也可手动修改'),
                ('税率', '请输入税率，如：3'),
                ('预付款', '输入数字，自动转换大写'),
                ('预付款大写', '自动转换，也可手动修改'),
                ('尾款', '输入数字，自动转换大写'),
                ('尾款大写', '自动转换，也可手动修改'),
                ('费用表格图片', '请粘贴费用表格图片路径'),
            ]),
            ("开票信息", [
                ('开票内容', '请输入开票内容'),
            ]),
            ("乙方财务信息", [
                ('乙方银行账号', '请输入银行账号'),
                ('乙方银行开户行', '请输入开户行名称'),
            ]),
            ("乙方联系人信息", [
                ('乙方代表名称', '请输入联系人姓名'),
                ('乙方代表电话', '请输入联系电话'),
                ('乙方代表邮箱', '请输入联系邮箱'),
            ]),
            ("乙方地址", [
                ('乙方地址', '请输入完整地址，将自动分行'),
            ]),
        ]

        # 为每个分组创建UI
        for group_name, fields in groups:
            group_box = QGroupBox(group_name)
            form = QFormLayout(group_box)

            for field_key, placeholder_text in fields:
                # 构建占位符 key（不跳过缺失的占位符——officecli 负责跨 run 替换）
                full_placeholder = f'替换的{field_key}' if not field_key.startswith('乙方银行') else field_key

                # 特殊处理地址字段
                if field_key == '乙方地址' and full_placeholder not in self.template_processor.placeholders:
                    line_edit = QLineEdit()
                    line_edit.setPlaceholderText(placeholder_text)
                    self.input_fields['乙方地址'] = line_edit
                    form.addRow("乙方地址:", line_edit)
                    continue

                # 创建标签
                label = f"{field_key}:"

                # 创建输入框（图片路径使用可拖拽的 DropLineEdit）
                if '图片' in field_key or '表格图片' in field_key:
                    line_edit = DropLineEdit()
                else:
                    line_edit = QLineEdit()
                line_edit.setPlaceholderText(placeholder_text)

                # 金额字段自动转换大写
                if ('费用' in field_key or '款' in field_key) and '大写' not in field_key and '表格' not in field_key:
                    line_edit.textChanged.connect(
                        lambda text, p=full_placeholder: self.on_amount_changed(p, text)
                    )

                self.input_fields[full_placeholder] = line_edit
                form.addRow(label, line_edit)

            self.form_layout.addWidget(group_box)

        self.form_layout.addStretch()

    def on_amount_changed(self, field_name: str, text: str):
        """金额变化时自动转换大写"""
        try:
            amount = float(text)
            chinese = amount_to_chinese(amount)

            # 查找对应的大写字段
            big_field_name = field_name + '大写'
            if big_field_name in self.input_fields:
                self.input_fields[big_field_name].setText(chinese)
        except ValueError:
            pass

    def refresh_contact_combo(self):
        """刷新联系人下拉框和列表"""
        names = self.contact_manager.get_names()
        self.contact_combo.clear()
        self.contact_combo.addItems(names)
        self.contact_list.clear()
        self.contact_list.addItems(names)

    def on_contact_list_clicked(self, item):
        """联系人列表点击事件"""
        name = item.text()
        contact = self.contact_manager.get_contact(name)
        if contact:
            self.contact_combo.setCurrentText(name)
            self.fill_form(contact)

    def on_contact_selected(self, name: str):
        """联系人选择变化"""
        pass

    def load_selected_contact(self):
        """加载选中的联系人信息"""
        name = self.contact_combo.currentText()
        if not name:
            return

        contact = self.contact_manager.get_contact(name)
        if contact:
            self.fill_form(contact)

    def fill_form(self, data: Dict[str, str]):
        """用数据填充表单"""
        for key, value in data.items():
            if key in self.input_fields:
                self.input_fields[key].setText(str(value))

    def save_current_as_contact(self):
        """将当前表单保存为联系人"""
        data = self.get_form_data()

        name = data.get('替换的乙方名称', '') or data.get('乙方名称', '')

        if not name:
            QMessageBox.warning(self, "警告", "请先填写乙方名称！")
            return

        existing = self.contact_manager.get_contact(name)
        if existing:
            reply = QMessageBox.question(
                self, "联系人已存在",
                f"联系人 '{name}' 已存在，是否覆盖？\n\n"
                "选择「是」覆盖已有联系人\n"
                "选择「否」取消保存",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply != QMessageBox.StandardButton.Yes:
                return

        self.contact_manager.add_contact(data)
        self.refresh_contact_combo()
        QMessageBox.information(self, "成功", f"联系人 '{name}' 已保存！")

    def clear_form_for_new_contact(self):
        """清空表单以创建新联系人"""
        for line_edit in self.input_fields.values():
            line_edit.clear()
        self.contact_combo.setCurrentIndex(-1)
        self.contact_list.clearSelection()

    def delete_selected_contact(self):
        """删除选中的联系人"""
        name = self.contact_combo.currentText()
        if not name:
            return

        reply = QMessageBox.question(
            self, "确认删除",
            f"确定要删除联系人 '{name}' 吗？",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.contact_manager.delete_contact(name)
            self.refresh_contact_combo()

    def select_output_dir(self):
        """选择输出目录"""
        dir_path = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if dir_path:
            self.output_dir.setText(dir_path)
            self._save_settings()

    def _load_settings(self):
        """加载上次的模板路径和输出目录"""
        try:
            if os.path.exists(self.settings_file):
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    settings = json.load(f)
                template = settings.get('template_path', '')
                if template and os.path.exists(template):
                    self.template_path.setText(template)
                    self.load_template()
                output_dir = settings.get('output_dir', '')
                if output_dir:
                    self.output_dir.setText(output_dir)
        except Exception:
            pass

    def _save_settings(self):
        """保存当前模板路径和输出目录"""
        try:
            os.makedirs(self.config_dir, exist_ok=True)
            settings = {
                'template_path': self.template_path.text(),
                'output_dir': self.output_dir.text(),
            }
            with open(self.settings_file, 'w', encoding='utf-8') as f:
                json.dump(settings, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def get_form_data(self) -> Dict[str, str]:
        """获取表单数据"""
        data = {}
        for key, line_edit in self.input_fields.items():
            data[key] = line_edit.text()

        # 处理地址字段：将"乙方地址"转换为三个地址行
        if '乙方地址' in data:
            address = data.pop('乙方地址')
            if address:
                data['乙方地址'] = address  # 保留原始地址，供generate方法处理

        return data

    def get_output_path(self, ext: str) -> str:
        """获取输出文件路径，自动重命名：合同编号_项目名称（乙方名称）"""
        out_dir = self.output_dir.text() or os.path.dirname(self.template_path.text())

        name = self.output_name.text()
        if not name:
            data = self.get_form_data()
            contract_no = data.get('替换的合同编号', '')
            project_name = data.get('替换的项目名称', '')
            company = data.get('替换的乙方名称', '')

            parts = []
            if contract_no:
                parts.append(contract_no)
            if project_name and company:
                parts.append(f"{project_name}（{company}）")
            elif project_name:
                parts.append(project_name)
            elif company:
                parts.append(company)

            name = '_'.join(parts) if parts else "合同"

        return os.path.join(out_dir, f"{name}.{ext}")

    def generate(self):
        """生成合同"""
        if not self.template_processor:
            QMessageBox.warning(self, "警告", "请先加载模板！")
            return

        data = self.get_form_data()

        empty_fields = [k for k, v in data.items() if not v.strip()]
        if empty_fields:
            reply = QMessageBox.question(
                self, "提示",
                f"以下字段为空:\n{chr(10).join(empty_fields[:5])}...\n\n确定要继续生成吗？",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply != QMessageBox.StandardButton.Yes:
                return

        try:
            docx_path = self.get_output_path('docx')
            self.template_processor.generate(data, docx_path)
            QMessageBox.information(self, "成功", f"DOCX已生成:\n{docx_path}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成失败:\n{str(e)}")


# ============================================================
# 程序入口
# ============================================================

def main():
    app = QApplication(sys.argv)

    # 设置应用样式
    app.setStyle('Fusion')

    # 设置字体
    font = QFont()
    font.setPointSize(10)
    app.setFont(font)

    # 深色主题样式表（深色底 + 紫色点缀）
    app.setStyleSheet("""
        /* 全局样式 */
        QWidget {
            background-color: #1e1e2e;
            color: #cdd6f4;
            font-family: "Microsoft YaHei", "Segoe UI", sans-serif;
        }

        /* 主窗口 */
        QMainWindow {
            background-color: #1e1e2e;
        }

        /* 分组框 */
        QGroupBox {
            background-color: #1e1e2e;
            border: 1px solid #45475a;
            margin-top: 14px;
            padding-top: 16px;
            font-weight: bold;
            color: #cba6f7;
        }

        QGroupBox::title {
            subcontrol-origin: margin;
            left: 12px;
            padding: 0 8px;
            color: #cba6f7;
        }

        /* 输入框 */
        QLineEdit {
            background-color: #313244;
            border: 1px solid #45475a;
            padding: 7px 10px;
            font-size: 13px;
            color: #cdd6f4;
            selection-background-color: #7c3aed;
            selection-color: #ffffff;
        }

        QLineEdit:focus {
            border-color: #cba6f7;
            background-color: #313244;
        }

        QLineEdit:hover {
            border-color: #585b70;
        }

        QLineEdit:read-only {
            background-color: #181825;
            color: #6c7086;
        }

        /* 按钮 - 默认灰色 */
        QPushButton {
            background-color: #45475a;
            color: #cdd6f4;
            border: none;
            padding: 7px 10px;
            font-size: 13px;
        }

        QPushButton:hover {
            background-color: #585b70;
        }

        QPushButton:pressed {
            background-color: #313244;
        }

        QPushButton:disabled {
            background-color: #313244;
            color: #6c7086;
        }

        /* 生成按钮 - 紫色 */
        QPushButton#btnGenDocx {
            background-color: #7c3aed;
            color: #ffffff;
            font-weight: bold;
            font-size: 15px;
            padding: 12px 24px;
        }

        QPushButton#btnGenDocx:hover {
            background-color: #6d28d9;
        }

        QPushButton#btnGenDocx:pressed {
            background-color: #5b21b6;
        }

        /* 下拉框 */
        QComboBox {
            background-color: #313244;
            border: 1px solid #45475a;
            padding: 7px 10px;
            font-size: 13px;
            color: #cdd6f4;
        }

        QComboBox:hover {
            border-color: #585b70;
        }

        QComboBox::drop-down {
            border: none;
            width: 28px;
        }

        QComboBox::down-arrow {
            image: none;
            border-left: 5px solid transparent;
            border-right: 5px solid transparent;
            border-top: 6px solid #cba6f7;
            margin-right: 10px;
        }

        QComboBox QAbstractItemView {
            background-color: #313244;
            border: 1px solid #cba6f7;
            selection-background-color: #7c3aed;
            selection-color: #ffffff;
            color: #cdd6f4;
        }

        /* 列表控件 */
        QListWidget {
            background-color: #313244;
            border: 1px solid #45475a;
            color: #cdd6f4;
            font-size: 13px;
        }

        QListWidget::item {
            padding: 6px 10px;
        }

        QListWidget::item:selected {
            background-color: #7c3aed;
            color: #ffffff;
        }

        QListWidget::item:hover {
            background-color: #45475a;
        }

        /* 标签 */
        QLabel {
            color: #cdd6f4;
            font-size: 13px;
        }

        QLabel#appTitle {
            font-size: 20px;
            font-weight: bold;
            color: #cba6f7;
        }

        QLabel#authorLabel {
            font-size: 12px;
            color: #6c7086;
        }

        /* 滚动区域 */
        QScrollArea {
            border: none;
            background-color: transparent;
        }

        /* 滚动条 */
        QScrollBar:vertical {
            background-color: #181825;
            width: 10px;
            margin: 0;
        }

        QScrollBar::handle:vertical {
            background-color: #45475a;
            min-height: 30px;
        }

        QScrollBar::handle:vertical:hover {
            background-color: #cba6f7;
        }

        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
            height: 0;
        }

        QScrollBar:horizontal {
            background-color: #181825;
            height: 10px;
            margin: 0;
        }

        QScrollBar::handle:horizontal {
            background-color: #45475a;
            min-width: 30px;
        }

        QScrollBar::handle:horizontal:hover {
            background-color: #cba6f7;
        }

        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
            width: 0;
        }

        /* 分割器 */
        QSplitter::handle {
            background-color: #45475a;
            width: 2px;
        }

        QSplitter::handle:hover {
            background-color: #cba6f7;
        }

        /* 消息框 */
        QMessageBox {
            background-color: #1e1e2e;
        }

        QMessageBox QLabel {
            color: #cdd6f4;
        }

        QMessageBox QPushButton {
            min-width: 80px;
        }

        /* 文件对话框 */
        QFileDialog {
            background-color: #1e1e2e;
        }
    """)

    # 创建主窗口
    window = MainWindow()
    window.show()

    sys.exit(app.exec())


if __name__ == '__main__':
    main()
